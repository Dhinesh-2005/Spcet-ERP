import sys
import os
import io
import re
from io import BytesIO
import docx
from docx.shared import Inches, Pt
import openpyxl
from PIL import Image, ImageDraw

# Set UTF-8 encoding for stdout
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

# Add current dir to path
sys.path.insert(0, os.path.dirname(__file__))

from app.utils.docx_parser import parse_docx_question_bank_rows, extract_cell_text
from app.routes.question_bank import process_question_equation, _normalize_co, _normalize_part, _normalize_unit, _extract_all_excel_images

def create_sample_image_bytes():
    img = Image.new("RGB", (100, 100), color=(73, 109, 137))
    d = ImageDraw.Draw(img)
    d.text((10, 10), "Test Img", fill=(255, 255, 0))
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()

def create_sample_docx():
    doc = docx.Document()
    doc.add_heading("Sample Question Bank Document", level=1)
    
    # Section 1 / Table 1
    doc.add_heading("Section A: Short Answer Questions", level=2)
    table1 = doc.add_table(rows=1, cols=7)
    hdr_cells = table1.rows[0].cells
    hdr_titles = ["S.No", "Question", "Image", "CO", "K-Level", "Mark", "Part"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        
    img_bytes = create_sample_image_bytes()
    img_stream = BytesIO(img_bytes)

    # Row 1: Math & Chemical equation
    r1 = table1.add_row().cells
    r1[0].text = "1"
    r1[1].text = "Balance the chemical equation H₂ + O₂ → H₂O and calculate the enthalpy."
    r1[2].text = "" # empty cell
    r1[3].text = "CO1"
    r1[4].text = "K2"
    r1[5].text = "2"
    r1[6].text = "A"

    # Row 2: Equation with exponent x² and integral ∫
    r2 = table1.add_row().cells
    r2[0].text = "2"
    r2[1].text = "Solve the equation f(x) = x² + 3x - 5 and find the value of integral ∫ x dx."
    r2[2].text = ""
    r2[3].text = "CO1"
    r2[4].text = "K3"
    r2[5].text = "2"
    r2[6].text = "A"

    # Row 3: Question with image
    r3 = table1.add_row().cells
    r3[0].text = "3"
    r3[1].text = "Identify the logic gate shown in the diagram below and state its truth table."
    # Insert image into cell
    p = r3[2].paragraphs[0]
    p.add_run().add_picture(img_stream, width=Inches(0.8))
    r3[3].text = "CO2"
    r3[4].text = "K2"
    r3[5].text = "2"
    r3[6].text = "A"

    # Row 4: Empty optional cells for CO, K-Level, Part
    r4 = table1.add_row().cells
    r4[0].text = "4"
    r4[1].text = "What is a finite automaton? Define deterministic finite automata (DFA)."
    r4[2].text = ""
    r4[3].text = "" # Empty CO
    r4[4].text = "" # Empty K-Level
    r4[5].text = "2"
    r4[6].text = "" # Empty Part

    # Section 2 / Table 2 (Multiple tables support test)
    doc.add_heading("Section B: Long Answer Questions", level=2)
    table2 = doc.add_table(rows=1, cols=7)
    hdr_cells2 = table2.rows[0].cells
    for i, title in enumerate(hdr_titles):
        hdr_cells2[i].text = title

    r5 = table2.add_row().cells
    r5[0].text = "5"
    r5[1].text = "Explain the architecture of 8086 microprocessor with a neat block diagram."
    r5[2].text = ""
    r5[3].text = "CO1"
    r5[4].text = "K3"
    r5[5].text = "16"
    r5[6].text = "B"

    r6 = table2.add_row().cells
    r6[0].text = "6"
    r6[1].text = "Derive the expression for the area of a circle using double integral ∫∫ r dr dθ."
    r6[2].text = ""
    r6[3].text = "CO1"
    r6[4].text = "K3"
    r6[5].text = "16"
    r6[6].text = "B"

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def create_sample_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Question Bank"

    hdr = ["S.No", "Question", "Image", "CO", "K-Level", "Mark", "Part"]
    ws.append(hdr)

    ws.append(["1", "Define OSI Reference Model and explain seven layers.", "", "CO1", "K1", "2", "A"])
    ws.append(["2", "What is TCP/IP protocol suite? Compare TCP vs UDP.", "", "CO1", "K2", "2", "A"])
    ws.append(["3", "Explain Dijkstra's shortest path algorithm with an example graph.", "", "CO1", "K3", "16", "B"])
    ws.append(["4", "Derive Bellman-Ford algorithm complexity for negative edge weights.", "", "CO1", "K3", "16", "B"])

    out = BytesIO()
    wb.save(out)
    return out.getvalue()

def main():
    print("=== STARTING DOCX & XLSX QUESTION BANK PARSER TESTS ===")

    # 1. Test DOCX parsing
    docx_bytes = create_sample_docx()
    print(f"\n1. Generated sample DOCX ({len(docx_bytes)} bytes)")
    docx_rows = parse_docx_question_bank_rows(docx_bytes)
    print(f"Extracted {len(docx_rows)} question rows from DOCX tables:")

    for idx, r in enumerate(docx_rows):
        has_img = "YES" if r.get("image") else "NO"
        q_preview = r['question'][:50].replace('\n', ' ')
        print(f"  [{idx+1}] Q: '{q_preview}...' | CO: '{r['co']}' | K: '{r['kLevel']}' | Mark: '{r['marks']}' | Part: '{r['part']}' | Has Image: {has_img}")

    assert len(docx_rows) == 6, f"Expected 6 rows from docx, got {len(docx_rows)}"
    assert docx_rows[2]["image"] is not None, "Expected embedded image in row 3"
    assert "H₂" in docx_rows[0]["question"], "Expected chemical formula in row 1"

    # 2. Test equation processor on extracted DOCX text
    eq_result = process_question_equation(docx_rows[0]["question"])
    print(f"\n2. Equation processor on row 1:\n   Question: {eq_result['question']}\n   LaTeX: {eq_result['latex']}\n   hasEquation: {eq_result['hasEquation']}")

    # 3. Test XLSX parsing compatibility
    xlsx_bytes = create_sample_xlsx()
    print(f"\n3. Generated sample XLSX ({len(xlsx_bytes)} bytes)")
    wb = openpyxl.load_workbook(BytesIO(xlsx_bytes), data_only=True)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    print(f"Extracted {len(rows)-1} question rows from Excel.")
    assert len(rows) == 5, f"Expected 5 total rows (header + 4 questions), got {len(rows)}"

    print("\n=== ALL DOCX & XLSX TESTS PASSED SUCCESSFULLY! ===")

if __name__ == "__main__":
    main()
